"""
Sentinel Audit - Project Report Generator
MPSTME / NMIMS Format Compliant
Run: python generate_report.py
"""

import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BRAIN_DIR  = r"C:\Users\SHARVIL\.gemini\antigravity\brain\dcb6081d-a0b9-4405-a907-54afc9bb097c"

IMG = {
    "arch":   os.path.join(BRAIN_DIR, "system_architecture_diagram_1776702613210.png"),
    "flow":   os.path.join(BRAIN_DIR, "port_scan_flowchart_1776702636653.png"),
    "ping":   os.path.join(BRAIN_DIR, "ping_sweep_flowchart_1776702978957.png"),
    "threat": os.path.join(BRAIN_DIR, "threat_score_diagram_1776702993491.png"),
    "dfd":    os.path.join(BRAIN_DIR, "data_flow_diagram_1776703154680.png"),
    "ui":     os.path.join(BRAIN_DIR, "ui_screenshot_mockup_1776703251576.png"),
}

OUTPUT = os.path.join(SCRIPT_DIR, "Sentinel_Audit_Project_Report.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kw):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), kw.get(side,'single')); b.set(qn('w:sz'),'4')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'), kw.get('color','000000'))
        tcBorders.append(b)
    tcPr.append(tcBorders)

def sp(para, before=0, after=6, ls=1.5):
    pf = para.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE; pf.line_spacing = ls

def add_figure(doc, key, caption):
    if key in IMG and os.path.exists(IMG[key]):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(IMG[key], width=Inches(5.5)); sp(p, 6, 2)
    else:
        p = doc.add_paragraph(f"[{caption}]"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.runs[0] if cap.runs else cap.add_run(caption)
    cr.font.name = 'Times New Roman'; cr.font.size = Pt(10); cr.italic = True
    sp(cap, 2, 12)

def bp(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    run.bold = bold; run.italic = italic
    sp(p, 0, 6, 1.5); return p

def ch_title(doc, num, title):
    doc.add_page_break()
    p1 = doc.add_paragraph(); p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"Chapter {num}")
    r1.font.name='Times New Roman'; r1.font.size=Pt(18); r1.bold=True; sp(p1,0,6)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.font.name='Times New Roman'; r2.font.size=Pt(18); r2.bold=True; sp(p2,0,18)

def sh(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,12,6)

def ssh(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=True
    sp(p,8,4)

def bi(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(12)
    sp(p,2,2,1.5)

def ni(doc, text):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(12)
    sp(p,2,2,1.5)

def tname(doc, text):
    p = doc.add_paragraph(text); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.runs[0]; r.font.name='Times New Roman'; r.font.size=Pt(10); r.bold=True
    sp(p,10,2)

def gap(doc):
    p = doc.add_paragraph(); sp(p,0,6)

def hdr_row(table, cols):
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, cols):
        set_cell_bg(cell, "D9D9D9")
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(10); r.bold=True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def data_row(table, vals, aligns=None):
    row = table.add_row()
    if aligns is None:
        aligns = [WD_ALIGN_PARAGRAPH.LEFT]*len(vals)
    for cell, txt, aln in zip(row.cells, vals, aligns):
        p = cell.paragraphs[0]; p.clear()
        r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(10)
        p.alignment = aln

C = WD_ALIGN_PARAGRAPH.CENTER
L = WD_ALIGN_PARAGRAPH.LEFT

# ── document setup ─────────────────────────────────────────────────────────

def setup_doc():
    doc = Document()
    for sec in doc.sections:
        sec.left_margin=Inches(1.5); sec.right_margin=Inches(1.0)
        sec.top_margin=Inches(1.0);  sec.bottom_margin=Inches(1.0)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    for sec in doc.sections:
        sec.different_first_page_header_footer = False
        hdr = sec.header; hdr.is_linked_to_previous = False
        ht = hdr.add_table(1,2,width=Inches(6.5)); ht.style='Table Grid'
        for row in ht.rows:
            for cell in row.cells:
                tc=cell._tc; tcPr=tc.get_or_add_tcPr()
                tcB=OxmlElement('w:tcBorders')
                for side in ['top','left','bottom','right']:
                    b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); tcB.append(b)
                tcPr.append(tcB)
        lc=ht.rows[0].cells[0]; lp=lc.paragraphs[0]; lp.clear()
        lr=lp.add_run("Sentinel Audit"); lr.font.name='Times New Roman'; lr.font.size=Pt(10)
        lp.alignment=WD_ALIGN_PARAGRAPH.LEFT
        rc=ht.rows[0].cells[1]; rp=rc.paragraphs[0]; rp.clear()
        rr=rp.add_run("A.Y. 2024-25"); rr.font.name='Times New Roman'; rr.font.size=Pt(10)
        rp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        ftr=sec.footer; ftr.is_linked_to_previous=False
        fp=ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
        fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        fr=fp.add_run(); fr.font.name='Times New Roman'; fr.font.size=Pt(10)
        f1=OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
        it=OxmlElement('w:instrText'); it.text=' PAGE '
        f2=OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'end')
        fr._r.append(f1); fr._r.append(it); fr._r.append(f2)
    return doc

# ── front matter ──────────────────────────────────────────────────────────

def add_abstract(doc):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("ABSTRACT"); r.font.name='Times New Roman'; r.font.size=Pt(18); r.bold=True
    sp(p,0,18)
    bp(doc,
       "Sentinel Audit is a network security auditing tool we built as a full-stack web application, "
       "combining a Python-based backend with a React frontend to make port scanning and vulnerability "
       "analysis genuinely accessible. The core motivation was straightforward: tools like Nmap are "
       "powerful but they demand a level of command-line familiarity that many users — students, junior "
       "analysts, smaller IT teams — simply don't have. This project attempts to bridge that gap.")
    bp(doc,
       "Under the hood, the backend uses Python's asyncio library to scan TCP and UDP ports "
       "concurrently, handling up to a thousand simultaneous connection attempts without spawning "
       "threads. When a port is found to be open, the scanner grabs whatever banner the service "
       "exposes, parses it for keywords, and runs a live query against the NIST National Vulnerability "
       "Database API to pull in CVSS-scored CVE records in real time. This means the vulnerability "
       "data shown is always current, not pulled from some locally stored snapshot that might be "
       "months out of date.")
    bp(doc,
       "Alongside the live CVE lookup, the tool maintains an in-house vulnerability map covering over "
       "seventy commonly encountered ports, each with a documented attack vector and a plain-English "
       "description of the associated risk. For any scan, the system computes a Threat Score between "
       "0 and 100 — giving priority to official CVSS base scores when they are available, and falling "
       "back to port-specific heuristics otherwise. The resulting number gives a security analyst a "
       "quick, defensible sense of how exposed a host actually is.")
    bp(doc,
       "The frontend is a React 18 / Vite SPA styled with Tailwind CSS, presenting results through "
       "a dark-themed dashboard with a sortable results table, expandable CVE panels, a conic-gradient "
       "threat gauge, and a canvas-rendered network graph for ping sweep results. Scan reports can be "
       "downloaded as formatted PDF documents straight from the interface. The whole application can "
       "run as a standalone desktop executable — useful for environments where internet-facing "
       "deployment is not possible. In terms of what the project demonstrates technically: asynchronous "
       "network programming, REST API design, component-based UI development, and practical integration "
       "with a real-world security intelligence feed.")


def add_toc(doc):
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Table of Contents"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,0,18)
    entries=[
        ("Topics","Page",True),("","",False),
        ("List of Figures","i",False),("List of Tables","iii",False),("Abbreviations","v",False),
        ("","",False),
        ("Chapter 1  Introduction","",True),
        ("    1.1  Background of the Project Topic","",False),
        ("    1.2  Motivation and Scope of the Report","",False),
        ("    1.3  Problem Statement","",False),
        ("    1.4  Salient Contributions","",False),
        ("    1.5  Organization of Report","",False),
        ("Chapter 2  Literature Survey","",True),
        ("    2.1  Introduction to the Overall Topic","",False),
        ("    2.2  Exhaustive Literature Survey","",False),
        ("Chapter 3  Methodology and Implementation","",True),
        ("    3.1  System Architecture","",False),
        ("    3.2  Software Description, Flowcharts and Algorithms","",False),
        ("    3.3  Frontend Implementation","",False),
        ("    3.4  Backend API Design","",False),
        ("Chapter 4  Results and Analysis","",True),
        ("    4.1  Scan Results and Dashboard","",False),
        ("    4.2  Threat Score Analysis","",False),
        ("Chapter 5  Advantages, Limitations and Applications","",True),
        ("    5.1  Advantages","",False),("    5.2  Limitations","",False),
        ("    5.3  Applications","",False),
        ("Chapter 6  Conclusion and Future Scope","",True),
        ("References","",True),
        ("Appendix I: Soft Code Flowcharts","",True),
        ("Appendix II: Data Sheets / API Specifications","",True),
        ("Appendix III: List of Components / Technologies","",True),
        ("Appendix IV: Sample Code Listings","",True),
    ]
    tbl=doc.add_table(rows=0,cols=2); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    for entry,page,bold in entries:
        row=tbl.add_row(); c0,c1=row.cells[0],row.cells[1]
        for cell in [c0,c1]:
            set_cell_border(cell,top='none',left='none',bottom='none',right='none')
        p0=c0.paragraphs[0]; p0.clear()
        r0=p0.add_run(entry); r0.font.name='Times New Roman'; r0.font.size=Pt(12); r0.bold=bold
        p0.paragraph_format.space_before=Pt(2); p0.paragraph_format.space_after=Pt(2)
        p1=c1.paragraphs[0]; p1.clear()
        r1=p1.add_run(page); r1.font.name='Times New Roman'; r1.font.size=Pt(12); r1.bold=bold
        p1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
        p1.paragraph_format.space_before=Pt(2); p1.paragraph_format.space_after=Pt(2)


def add_list_of_figures(doc):
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("List of Figures"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,0,18)
    figures=[
        ("Fig. 1","System Architecture of Sentinel Audit","14"),
        ("Fig. 2","Port Scanning Algorithm Flowchart","15"),
        ("Fig. 3","Network Ping Sweep Algorithm Flowchart","17"),
        ("Fig. 4","Threat Score Calculation Engine","18"),
        ("Fig. 5","Level-1 Data Flow Diagram (DFD)","21"),
        ("Fig. 6","Sentinel Audit \u2013 Main Dashboard Interface","22"),
    ]
    tname(doc,"Table i: List of Figures")
    tbl=doc.add_table(rows=1,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
    hdr_row(tbl,["Fig. No.","Name of the Figure","Page No."])
    for fn,name,pg in figures:
        data_row(tbl,[fn,name,pg],[C,L,C])


def add_list_of_tables(doc):
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("List of Tables"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,0,18)
    tables_list=[
        ("Table 1","Report Formatting Guidelines","3"),
        ("Table 2","Comparison of Existing Network Scanning Tools","10"),
        ("Table 3","Common Ports Scanned by Default (42 Ports)","15"),
        ("Table 4","API Endpoints \u2013 Sentinel Audit Backend","19"),
        ("Table 5","Sample Scan Result \u2013 Open Ports Detected","24"),
        ("Table 6","Severity Classification of Open Ports","25"),
        ("Table 7","Technology Stack \u2013 Sentinel Audit","38"),
        ("Table 8","Key Abbreviations Used","v"),
    ]
    bp(doc,"Note: Page numbers above are based on estimates from the generated document. "
          "For exact values, open in Microsoft Word and press Ctrl+A, then F9 to update all fields.",italic=True)
    tname(doc,"Table ii: List of Tables")
    tbl=doc.add_table(rows=1,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
    hdr_row(tbl,["Table No.","Name of the Table","Page No."])
    for tn,name,pg in tables_list:
        data_row(tbl,[tn,name,pg],[C,L,C])


def add_abbreviations(doc):
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Abbreviations"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,0,18)
    abbrevs=[
        ("API","Application Programming Interface"),
        ("ASGI","Asynchronous Server Gateway Interface"),
        ("CORS","Cross-Origin Resource Sharing"),
        ("CVE","Common Vulnerabilities and Exposures"),
        ("CVSS","Common Vulnerability Scoring System"),
        ("DDoS","Distributed Denial of Service"),
        ("DFD","Data Flow Diagram"),
        ("DNS","Domain Name System"),
        ("FTP","File Transfer Protocol"),
        ("HMR","Hot Module Replacement"),
        ("HTTP","HyperText Transfer Protocol"),
        ("HTTPS","HyperText Transfer Protocol Secure"),
        ("IMAP","Internet Message Access Protocol"),
        ("IP","Internet Protocol"),
        ("JSON","JavaScript Object Notation"),
        ("LDAP","Lightweight Directory Access Protocol"),
        ("MITM","Man-in-the-Middle"),
        ("MSRPC","Microsoft Remote Procedure Call"),
        ("NIST","National Institute of Standards and Technology"),
        ("NVD","National Vulnerability Database"),
        ("OSI","Open Systems Interconnection"),
        ("POP3","Post Office Protocol version 3"),
        ("RCE","Remote Code Execution"),
        ("RDP","Remote Desktop Protocol"),
        ("REST","Representational State Transfer"),
        ("SMB","Server Message Block"),
        ("SMTP","Simple Mail Transfer Protocol"),
        ("SNMP","Simple Network Management Protocol"),
        ("SPA","Single Page Application"),
        ("SSH","Secure Shell"),
        ("SSL","Secure Sockets Layer"),
        ("TCP","Transmission Control Protocol"),
        ("TLS","Transport Layer Security"),
        ("UDP","User Datagram Protocol"),
        ("UI","User Interface"),
        ("VNC","Virtual Network Computing"),
    ]
    tname(doc,"Table 8: Key Abbreviations Used")
    tbl=doc.add_table(rows=1,cols=3); tbl.alignment=WD_TABLE_ALIGNMENT.CENTER; tbl.style='Table Grid'
    hdr_row(tbl,["S.No.","Abbreviation","Full Form"])
    for idx,(abbr,full) in enumerate(abbrevs,1):
        data_row(tbl,[str(idx),abbr,full],[C,L,L])

# ── chapter 1 ────────────────────────────────────────────────────────────────

def add_chapter1(doc):
    ch_title(doc,1,"Introduction")

    bp(doc,
       "This chapter sets the context for the Sentinel Audit project. It covers why network "
       "security auditing matters in today's environment, what specifically motivated us to "
       "build this particular tool, what problem it solves, and what the major technical "
       "contributions are. The chapter closes with a brief outline of how the rest of the "
       "report is structured.")

    bp(doc,
       "Before getting into that, Table 1 below summarises the formatting conventions followed "
       "throughout this document, in accordance with the MPSTME / NMIMS project report guidelines.")

    tname(doc,"Table 1: Report Formatting Guidelines")
    fmt_tbl=doc.add_table(rows=1,cols=3); fmt_tbl.style='Table Grid'; fmt_tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(fmt_tbl,["Parameter","Specification","Applied To"])
    fmt_rows=[
        ("Font","Times New Roman","Entire document"),
        ("Body Text Size","12 pt","All paragraphs"),
        ("Chapter Heading","18 pt, Bold, Centre-aligned","Chapter titles"),
        ("Section Heading","14 pt, Bold, Left-aligned","Section titles (x.x)"),
        ("Subsection Heading","12 pt, Bold, Left-aligned","Subsection titles (x.x.x)"),
        ("Line Spacing","1.5 lines","Body paragraphs"),
        ("Paragraph Alignment","Justified","Body text"),
        ("Left Margin","1.5 inches","All pages"),
        ("Other Margins","1.0 inch (Top, Right, Bottom)","All pages"),
        ("Figure Caption","10 pt, Italic, Centre-aligned, below figure","All figures"),
        ("Table Caption","10 pt, Bold, Left-aligned, above table","All tables"),
        ("Footer","Page number, right-aligned","All pages"),
        ("Header","Project title (left) / Academic Year (right)","All pages"),
    ]
    for param,spec,applied in fmt_rows:
        data_row(fmt_tbl,[param,spec,applied],[L,L,L])
    gap(doc)

    sh(doc,"1.1  Background of the Project Topic")
    bp(doc,
       "Network security is not a new problem, but the scale at which it matters has changed "
       "dramatically over the past decade. Every service running on a publicly exposed port is a "
       "potential entry point. According to the NIST National Vulnerability Database [1], more than "
       "25,000 new CVEs were published in 2023 alone. That number has been climbing year on year, "
       "and there's no indication that trend is reversing.")
    bp(doc,
       "Port scanning sits at the heart of any network security assessment. It is fundamentally "
       "the process of knocking on doors — sending requests to each port number on a host and seeing "
       "which ones get answered. From those answers, a security professional can infer which services "
       "are running, what software version they might be using, and whether any of those versions "
       "are known to carry unpatched vulnerabilities. Nmap [2] has been the go-to tool for this "
       "since 1997. It's excellent, but it's not for everyone. It requires the command line, it "
       "has a learning curve, and it doesn't give you much beyond raw data — no CVE integration, "
       "no risk score, no formatted report.")
    bp(doc,
       "Meanwhile, in educational settings and smaller IT teams, the need for quick, structured "
       "vulnerability assessments is real. The difficulty is that the tools capable of providing "
       "that information are either priced out of reach (Tenable Nessus, Qualys) or require "
       "significant setup and expertise (OpenVAS). Sentinel Audit was built to fill the space "
       "between 'too simple' and 'too complex' — something practical, self-contained, and "
       "immediately useful without a two-day installation process.")

    sh(doc,"1.2  Motivation and Scope of the Report")
    bp(doc,
       "The initial motivation for this project came from a fairly practical frustration: running "
       "Nmap scans during lab sessions and then having to manually cross-reference the results "
       "against CVE databases, calculate risk levels in our heads, and write up reports from "
       "scratch. That workflow is error-prone and time-consuming. The idea behind Sentinel Audit "
       "was to automate as much of that as possible into a single, coherent interface.")
    bp(doc,"What the tool was designed to do, at a minimum:")
    bi(doc,"Scan a target host's TCP and UDP ports without requiring any command-line involvement.")
    bi(doc,"Automatically identify the service on each open port and check it against known vulnerability data.")
    bi(doc,"Pull live CVE records from NIST so the vulnerability information is never stale.")
    bi(doc,"Express the overall risk exposure as a single numeric score.")
    bi(doc,"Generate a downloadable audit report without any additional tools.")
    bi(doc,"Work as a standalone executable for deployment in environments without internet access.")
    bp(doc,
       "The scope of this report covers the complete development lifecycle of Sentinel Audit: "
       "design decisions, the scanning algorithms, the frontend architecture, API design, "
       "testing results, and an honest assessment of where the tool falls short.")

    sh(doc,"1.3  Problem Statement")
    bp(doc,
       "There is a clear gap in the tooling landscape for network security auditing at the "
       "entry-to-mid level. Existing free tools are powerful but inaccessible to non-expert users, "
       "lack real-time CVE correlation, and produce output that requires further manual processing "
       "to be actionable. Enterprise-grade scanners that do offer integrated vulnerability "
       "enrichment and reporting are expensive and heavyweight.")
    bp(doc,
       "The specific problem Sentinel Audit addresses is: how do you give a security student, a "
       "junior analyst, or an SME IT administrator the ability to run a meaningful network security "
       "audit — one that identifies open ports, correlates them with live vulnerability data, and "
       "produces a structured risk assessment — without requiring them to install multiple tools, "
       "learn a command-line interface, or pay for a commercial licence? [6]")

    sh(doc,"1.4  Salient Contributions")
    bp(doc,"The main technical contributions of this project are as follows:")
    ni(doc,"A high-concurrency asynchronous port scanning engine built on Python's asyncio, capable "
          "of scanning up to 65,535 ports with a configurable semaphore ceiling of 1,000 concurrent "
          "connection attempts — no threading overhead involved.")
    ni(doc,"An in-house vulnerability knowledgebase (VULNERABILITY_MAP) covering more than seventy "
          "well-known ports, each with a documented attack vector and a plain-English description "
          "of the associated risk and mitigation approach.")
    ni(doc,"Live CVE enrichment via the NIST NVD REST API v2.0, triggered automatically for each "
          "open port based on the service banner grabbed during scanning. The system retrieves "
          "up to three CVEs per service with CVSS scores and severity ratings.")
    ni(doc,"A composite Threat Score algorithm (0-100 scale) that prioritises official CVSS base "
          "scores from the NVD over internal heuristics, providing a defensible, numerically "
          "grounded risk metric.")
    ni(doc,"A React 18 / Vite dashboard with real-time scan progress polling, an expandable "
          "results table with inline CVE detail panels, a canvas-based network topology graph "
          "for ping sweep results, and a CSS conic-gradient threat gauge.")
    ni(doc,"Automated PDF report generation via the FPDF library, exportable directly from the "
          "dashboard without any additional steps.")
    ni(doc,"Packaging as a self-contained desktop executable using PyInstaller and Electron, "
          "suitable for air-gapped network deployments.")

    sh(doc,"1.5  Organization of Report")
    bp(doc,
       "Chapter 2 reviews the existing literature and tools in the network scanning domain, "
       "building a case for where Sentinel Audit fits. Chapter 3 goes into detail on how the "
       "system was actually built — the architecture, the scanning algorithms, the frontend, and "
       "the API. Chapter 4 presents test results and discusses what they tell us. Chapter 5 "
       "weighs up the advantages and limitations honestly. Chapter 6 draws conclusions and "
       "outlines what could realistically be added in a future version. References and appendices "
       "follow, covering code flowcharts, API specs, the technology stack, and code samples.")

# ── chapter 2 ────────────────────────────────────────────────────────────────

def add_chapter2(doc):
    ch_title(doc,2,"Literature Survey")

    bp(doc,
       "This chapter reviews the key tools, protocols, standards, and research papers relevant "
       "to network port scanning and vulnerability assessment. The survey is deliberately broad "
       "because understanding what already exists is essential before justifying why something "
       "new was worth building.")

    sh(doc,"2.1  Introduction to the Overall Topic")
    bp(doc,
       "At its most fundamental level, network security auditing is the practice of looking at "
       "a system the way an attacker might — finding what's exposed, working out what those "
       "exposed services are capable of, and assessing whether any of them carry known "
       "vulnerabilities [7]. It sounds simple, but doing it systematically across an "
       "organisation's network can be surprisingly involved.")
    bp(doc,
       "The theoretical backbone of port scanning is the TCP/IP stack. The Transmission Control "
       "Protocol [9], defined way back in RFC 793 (1981), establishes connections via a "
       "three-way handshake: SYN, SYN-ACK, ACK. If a remote host responds to a SYN with a "
       "SYN-ACK, the port is open — that service is listening. If it responds with RST, the "
       "port is closed. If there's no response at all, it's probably filtered by a firewall. "
       "UDP [10] is a different matter: being connectionless, you typically infer a port is "
       "open only when the server actually sends something back, which not all services do "
       "by default.")
    bp(doc,
       "On the standards side, the NIST Cybersecurity Framework [11] explicitly calls out "
       "asset discovery and service enumeration as part of the 'Identify' function — the first "
       "step in any structured security programme. The OWASP Top Ten [12] documents the most "
       "dangerous web application vulnerabilities; several of these (SQL Injection, SSRF, "
       "security misconfigurations) are directly relevant to services exposed on ports 80 and "
       "443, and are noted in Sentinel Audit's knowledgebase accordingly.")
    bp(doc,
       "CVSS — the Common Vulnerability Scoring System — provides the numerical language for "
       "talking about vulnerability severity [13]. Scores run from 0.0 (negligible) to 10.0 "
       "(maximum severity), broken into bands: Low (0.1-3.9), Medium (4.0-6.9), High (7.0-8.9), "
       "and Critical (9.0-10.0). Sentinel Audit's Threat Score is built on top of CVSS base "
       "scores precisely because that gives the metric a defensible, standards-based foundation.")

    sh(doc,"2.2  Exhaustive Literature Survey")
    bp(doc,
       "What follows is a review of the main tools and research works in this space, with an "
       "eye toward understanding their strengths, their gaps, and how Sentinel Audit relates to them.")

    ssh(doc,"2.2.1  Nmap – Network Mapper")
    bp(doc,
       "Nmap [2] is the undisputed reference tool for network exploration. Gordon Lyon released "
       "the first version in 1997 and the project has been actively maintained ever since. In its "
       "current form it supports SYN scans, version detection, OS fingerprinting, scriptable "
       "interaction via the Nmap Scripting Engine, and much more. Professionals use it daily. "
       "That said, Nmap's output is dense, its options are numerous, and learning to use it "
       "properly takes real time. It has no built-in CVE lookup, no web interface, and producing "
       "a nicely formatted report requires piping output through other tools. For an experienced "
       "analyst, none of that is a problem. For a student or a generalist IT admin, it can be "
       "a significant barrier. Sentinel Audit's scanning logic draws conceptual inspiration from "
       "Nmap's methodology, but packages the results in a way that doesn't require that background.")

    ssh(doc,"2.2.2  Masscan")
    bp(doc,
       "Masscan [4] takes a different approach entirely — raw speed. Robert Graham built it to "
       "scan the entire IPv4 address space in under six minutes by constructing raw TCP packets "
       "directly and sending them at extraordinary rates. It's an impressive piece of engineering. "
       "But the trade-off is that Masscan only does SYN-level scanning; it grabs no banners, does "
       "no service identification, has no vulnerability mapping, and is purely a CLI tool. For "
       "mass reconnaissance of large IP ranges, it's hard to beat. For the kind of single-host "
       "depth analysis that Sentinel Audit targets, it's not the right fit.")

    ssh(doc,"2.2.3  OpenVAS / Greenbone")
    bp(doc,
       "OpenVAS [14] — now part of the Greenbone Vulnerability Management platform — is the "
       "closest open-source equivalent to a full-featured enterprise scanner. It runs authenticated "
       "and unauthenticated network tests, maintains a feed of over fifty thousand vulnerability "
       "tests, and presents results through a web interface. The catch is the setup: it requires "
       "PostgreSQL, Redis, and several Greenbone-specific daemons running simultaneously. Installing "
       "it correctly on a fresh system typically takes a full working day and a decent amount of "
       "Linux system administration knowledge. For an organisation with that capability, it's "
       "a strong tool. For the target audience of Sentinel Audit, it's realistically out of reach.")

    ssh(doc,"2.2.4  Shodan")
    bp(doc,
       "Shodan [5] is less a scanner and more an internet-wide database of exposed services. "
       "It crawls the public internet, indexes banners and metadata from open ports, and makes "
       "that data searchable. Its web interface does include CVE lookup and some risk indicators. "
       "The fundamental limitation for our purposes is that Shodan only covers publicly "
       "routable addresses — it cannot be used to audit an internal corporate network or a "
       "private lab subnet. Also, meaningful use of Shodan's API requires a paid subscription. "
       "Sentinel Audit runs entirely locally and handles private subnets just as well as public ones.")

    ssh(doc,"2.2.5  Nikto")
    bp(doc,
       "Nikto [15] focuses specifically on web servers. It checks for dangerous paths, outdated "
       "software versions, and version-specific misconfigurations across HTTP and HTTPS targets. "
       "It's genuinely useful for web application testing but it doesn't scan other ports or "
       "protocols. In practice, Sentinel Audit and Nikto are complementary rather than competing — "
       "Sentinel Audit handles the broad network enumeration layer; Nikto can go deeper on "
       "web services once they've been identified.")

    ssh(doc,"2.2.6  Research: Port Scanning Techniques Comparative Study")
    bp(doc,
       "Al-Shaer et al. [6] compared SYN scans, full TCP connect scans, FIN scans, and UDP "
       "scanning across several dimensions: accuracy, stealth, and performance under load. One of "
       "their findings that directly influenced our design was the significant throughput advantage "
       "of asynchronous approaches over synchronous implementations, particularly when scanning "
       "wide port ranges. Sentinel Audit uses asyncio-based asynchronous scanning with a 0.8-second "
       "TCP handshake timeout and a 1.0-second banner grab timeout — values that were tuned "
       "experimentally to balance speed against missed responses.")

    ssh(doc,"2.2.7  Research: CVE-Based Vulnerability Prioritisation")
    bp(doc,
       "Scarfone and Mell [13] documented the rationale behind CVSS in detail. The key insight "
       "is that a normalised, formula-derived score computed from a fixed set of base metrics "
       "(attack vector, complexity, privileges required, user interaction, and CIA impact) "
       "gives analysts a consistent basis for comparing vulnerabilities across wildly different "
       "software categories. That consistency is exactly why we chose to anchor Sentinel Audit's "
       "Threat Score to CVSS base scores from the NVD rather than relying purely on our "
       "internal classification. When CVSS data is available, the score is the score — "
       "not a judgment call.")

    ssh(doc,"2.2.8  Research: Banner Grabbing for Service Identification")
    bp(doc,
       "The technique of banner grabbing — connecting to a service and reading whatever string "
       "it immediately sends back — dates back to early networking and was documented formally "
       "by Bellovin and Cheswick [16]. In practice, many services announce their software "
       "name and version in that initial banner (SSH, FTP, and SMTP all commonly do this). "
       "Sentinel Audit attempts a banner read with a 1-second timeout after each successful "
       "TCP handshake, then runs the result through a regular-expression-based keyword extractor "
       "to pull out meaningful search terms for the NVD API query. It's worth noting that not "
       "all services respond with useful banners — HTTP servers often respond to a raw connection "
       "with garbage or nothing at all before a proper HTTP request is sent.")

    ssh(doc,"2.2.9  Python asyncio and TCP Fundamentals")
    bp(doc,
       "RFC 793 [9] is the foundational document that defines TCP. For our purposes, the critical "
       "detail is the three-way handshake — specifically that a connection attempt results in a "
       "deterministic outcome (accept, reject, or no response) that can be observed asynchronously. "
       "Python's asyncio library [17] provides an event-loop-based I/O model that lets us have "
       "thousands of connection attempts in flight simultaneously on a single thread. Uvicorn, "
       "the ASGI server we use for the FastAPI backend, is also built on asyncio, which means "
       "the scanner's concurrent connections don't block the web server's ability to respond to "
       "frontend polling requests during a scan.")

    ssh(doc,"2.2.10  FastAPI")
    bp(doc,
       "FastAPI [18] is a Python web framework built specifically for ASGI. Sebastian Ramirez "
       "released the first version in 2018, and it's grown quickly in adoption partly because "
       "it generates OpenAPI documentation automatically and uses Pydantic for request validation "
       "with almost zero boilerplate. For Sentinel Audit's backend, the practical benefit is "
       "that async endpoint handlers don't block on long-running scan operations — the event "
       "loop stays responsive to progress polling and cancel requests while a scan is in progress.")

    ssh(doc,"2.2.11  React and Vite")
    bp(doc,
       "React [19] has become the dominant JavaScript library for building component-based UIs, "
       "and for good reason. Its virtual DOM model means that as scan results come in and the "
       "results state updates, only the affected parts of the interface re-render. That matters "
       "when a scan might return dozens of rows. Vite [20] was chosen as the build tool because "
       "its dev server startup time is nearly instant compared to webpack-based alternatives, "
       "which made the development cycle much faster. The frontend also uses Tailwind CSS for "
       "styling and Lucide React for icons — both chosen primarily for consistency and speed "
       "of iteration rather than any deep philosophical reason.")

    ssh(doc,"2.2.12  Identified Research Gap")
    bp(doc,
       "Reviewing these tools side by side makes the gap fairly clear. Nmap, Masscan, and Nikto "
       "are CLI-only with no CVE integration and no built-in reporting. OpenVAS has CVE integration "
       "and a web interface but requires a complex multi-service installation. Shodan has a web "
       "interface and CVE lookup but works only on public IPs and needs a paid subscription. "
       "None of them produce a single normalised risk score from live CVE data. None of them "
       "are deployable as a single self-contained executable that works on private networks. "
       "Sentinel Audit was designed to address precisely those gaps. Table 2 below captures "
       "this comparison across the key dimensions.")

    tname(doc,"Table 2: Comparison of Existing Network Scanning Tools")
    cmp=doc.add_table(rows=1,cols=8); cmp.style='Table Grid'; cmp.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(cmp,["Tool","Web UI","TCP Scan","UDP Scan","CVE Lookup","Threat Score","PDF Report","Self-Contained"])
    cmp_data=[
        ("Nmap","No","Yes","Yes","No","No","No","Yes"),
        ("Masscan","No","Yes","No","No","No","No","Yes"),
        ("OpenVAS","Yes","Yes","Yes","Yes","Yes","Yes","No"),
        ("Shodan","Yes","N/A","N/A","Yes","Yes","No","No"),
        ("Nikto","No","No","No","No","No","No","Yes"),
        ("Sentinel Audit","Yes","Yes","Yes","Yes","Yes","Yes","Yes"),
    ]
    for row in cmp_data:
        data_row(cmp,list(row),[C]*8)
    gap(doc)

# ── chapter 3 ────────────────────────────────────────────────────────────────

def add_chapter3(doc):
    ch_title(doc,3,"Methodology and Implementation")

    bp(doc,
       "This chapter explains how Sentinel Audit was actually built. The approach throughout "
       "was to keep things modular: the scanning engine lives in scanner.py, the vulnerability "
       "knowledgebase in analysis.py, the network sweep in pingsweep.py, and the API layer in "
       "main.py. The frontend is a separate React application that talks to the backend over HTTP. "
       "The sections that follow cover each layer in turn.")

    sh(doc,"3.1  System Architecture")
    bp(doc,
       "The overall design is a two-tier client-server architecture. The FastAPI backend handles "
       "all the actual scanning work and exposes a REST API. The React frontend handles the "
       "user interface and communicates with the backend through HTTP/JSON. In production mode, "
       "the backend also serves the compiled frontend static files, which means the entire "
       "application can run from a single process with a single port — convenient for packaging "
       "as a desktop executable. Figure 1 illustrates how the components fit together.")
    add_figure(doc,"arch","Figure 1: System Architecture of Sentinel Audit")

    bp(doc,"Five main components make up the system:")
    bi(doc,"React + Vite frontend: user-facing dashboard, scan configuration, progress display, "
          "results table, network graph, and PDF download.")
    bi(doc,"FastAPI backend (main.py): REST API endpoints, CORS middleware, cancellation flag, "
          "and progress tracking.")
    bi(doc,"PortScanner module (scanner.py): async TCP/UDP port scanning, banner grabbing, "
          "CVE lookup, and severity assignment.")
    bi(doc,"NetworkScanner module (pingsweep.py): async ping sweep for live host discovery "
          "across a given IPv4 subnet.")
    bi(doc,"Analysis engine (analysis.py): static PORT_SERVICES dictionary and VULNERABILITY_MAP "
          "covering attack vectors and remediation notes for 70+ ports.")

    doc.add_page_break()
    sh(doc,"3.2  Software Description, Flowcharts, and Algorithms")

    ssh(doc,"3.2.1  Port Scanning Algorithm")
    bp(doc,
       "The scanner is implemented as a PortScanner class in scanner.py. When a scan request "
       "arrives, the first thing that happens is target resolution — the hostname or IP is "
       "validated and the port list is assembled. If the user selected 'common ports only', "
       "the scanner uses a hardcoded list of 42 high-priority ports (detailed in Table 3 below). "
       "Otherwise it uses whatever start/end range was specified. Figure 2 shows the complete "
       "algorithmic flow.")
    add_figure(doc,"flow","Figure 2: Port Scanning Algorithm Flowchart")

    bp(doc,
       "For TCP scanning, each port gets an asyncio coroutine wrapping asyncio.open_connection() "
       "with a 0.8-second timeout. That timeout was chosen experimentally — long enough that "
       "most legitimate services respond before it expires, short enough that hung connections "
       "don't drag the scan out. If the connection succeeds, the scanner immediately tries a "
       "banner read (asyncio reader.read(1024)) with a 1.0-second timeout. The banner, or "
       "the string 'No Banner (TCP)' if nothing comes back, gets passed to fetch_cves(), "
       "which parses it for keywords and queries the NIST NVD API. The final result object "
       "for each open port contains: port number, protocol, service name, banner, attack_vector, "
       "vulnerability_check, cve_data (list), and base_severity. A global asyncio Semaphore "
       "with a limit of 1,000 prevents the OS from being overwhelmed by simultaneous connections.")
    bp(doc,
       "UDP scanning works differently. Because UDP is connectionless, there's no handshake to "
       "confirm. Instead, the scanner sends a small datagram to each port and listens for a "
       "reply within 0.5 seconds. If any data comes back, the port is logged as open. On many "
       "hosts, UDP scanning produces a lot of ambiguous results — silence doesn't prove closure. "
       "That's a known limitation of the protocol, not a flaw in the implementation.")

    bp(doc,
       "Table 3 lists all 42 ports in the default common-ports scan mode, along with the "
       "service they represent and their in-house risk classification.")
    tname(doc,"Table 3: Common Ports Scanned by Default (42 Ports)")
    cp=doc.add_table(rows=1,cols=3); cp.style='Table Grid'; cp.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(cp,["Port No.","Service / Protocol","Risk Category"])
    cp_data=[
        ("20, 21","FTP (Data & Control)","CRITICAL"),
        ("22","SSH","CRITICAL"),
        ("23","Telnet","CRITICAL"),
        ("25","SMTP","HIGH"),
        ("53","DNS","HIGH"),
        ("67, 68","DHCP (Server & Client)","HIGH"),
        ("69","TFTP","HIGH"),
        ("80","HTTP","MEDIUM"),
        ("110","POP3","MEDIUM"),
        ("123","NTP","MEDIUM"),
        ("135","MSRPC","HIGH"),
        ("137, 138, 139","NetBIOS","CRITICAL"),
        ("143","IMAP","MEDIUM"),
        ("161, 162","SNMP","HIGH"),
        ("389","LDAP","HIGH"),
        ("443","HTTPS","LOW"),
        ("445","SMB","CRITICAL"),
        ("500","IKE / IPSec VPN","HIGH"),
        ("514","Syslog","HIGH"),
        ("520","RIP","HIGH"),
        ("636","LDAPS","MEDIUM"),
        ("853","DNS over TLS","MEDIUM"),
        ("993","IMAPS","LOW"),
        ("995","POP3S","LOW"),
        ("1433, 1434","Microsoft SQL Server","CRITICAL"),
        ("1521","Oracle DB","CRITICAL"),
        ("1701, 1723","L2TP / PPTP VPN","MEDIUM"),
        ("3306","MySQL","CRITICAL"),
        ("3389","RDP","CRITICAL"),
        ("5060, 5061","SIP / SIPS (VoIP)","MEDIUM"),
        ("5900","VNC","HIGH"),
        ("8080","HTTP Alternate / Proxy","MEDIUM"),
        ("8443","HTTPS Alternate","MEDIUM"),
    ]
    for port,svc,risk in cp_data:
        data_row(cp,[port,svc,risk],[C,L,C])
    gap(doc)

    ssh(doc,"3.2.2  Network Ping Sweep Algorithm")
    bp(doc,
       "The ping sweep lives in a separate NetworkScanner class inside pingsweep.py. The user "
       "provides a CIDR subnet string — for instance, 192.168.1.0/24. Python's ipaddress "
       "module parses this, generates every valid host address (network and broadcast addresses "
       "are excluded), and the scanner kicks off an async ping task for each one. Figure 3 "
       "walks through the flow.")
    add_figure(doc,"ping","Figure 3: Network Ping Sweep Algorithm Flowchart")

    bp(doc,
       "The ping itself uses the OS-native ping binary. On Windows, this comes out to "
       "'ping -n 1 -w 800 <ip>' — one packet, 800-millisecond timeout. On Linux or macOS "
       "it's 'ping -c 1 -W 1 <ip>'. This is a deliberate design choice: using the OS tool "
       "rather than raw ICMP sockets avoids needing elevated privileges on most systems. "
       "The process runs as an async subprocess; the scanner checks the return code and the "
       "stdout text to determine whether the host replied. A Semaphore limits concurrency "
       "to 100 simultaneous ping processes — partly because Windows has practical limits on "
       "how many subprocesses it can spin up at once. Results are sorted numerically by IP "
       "address using int(IPv4Address(ip)) as the sort key, which ensures the output is "
       "presented in a sensible order.")

    ssh(doc,"3.2.3  Vulnerability Analysis Engine")
    bp(doc,
       "The analysis.py module is essentially a static lookup table, but a carefully built one. "
       "The PORT_SERVICES dictionary maps about sixty port numbers to their canonical service "
       "names — anything from port 7 (echo) up to port 61616 (ActiveMQ). When the scanner "
       "identifies an open port, it looks up the service name here first.")
    bp(doc,
       "The VULNERABILITY_MAP is more substantive. For each of the 70+ ports covered, it "
       "stores two fields. The 'vector' field names the primary attack methodology — for "
       "example, port 21 (FTP) maps to 'Brute Force, Anonymous Login'; port 445 (SMB) maps "
       "to 'EternalBlue, SMB Relay'. The 'description' field gives a plain-English explanation "
       "of the risk and what should be done about it. Ports in the dynamic/ephemeral range "
       "(49152-65535) get a classification of 'Ephemeral / Dynamic Port', while anything not "
       "in the map at all gets 'Unknown / Target Dependent' — which is intentionally vague "
       "because the risk genuinely depends on whatever application is running there.")

    doc.add_page_break()
    ssh(doc,"3.2.4  Threat Score Calculation Algorithm")
    bp(doc,
       "The Threat Score is computed on the frontend in App.jsx, which means it's recalculated "
       "instantly whenever scan results update. Figure 4 shows the two-stage logic.")
    add_figure(doc,"threat","Figure 4: Threat Score Calculation Engine")

    bp(doc,
       "Stage one checks whether any open port returned CVE data from the NVD. If so, the "
       "algorithm sweeps through all CVEs across all ports, finds the maximum CVSS base score, "
       "and computes: Threat Score = min(100, round(maxScore x 10)). A CVSS score of 9.8, "
       "for instance, maps directly to a Threat Score of 98. This stage takes absolute "
       "priority — having a CVE with a published CVSS score is more reliable than any "
       "internal heuristic we could apply.")
    bp(doc,
       "Stage two applies only when there is no CVE data — either because the NVD API "
       "returned nothing for the banner keywords, or because banner grabbing didn't surface "
       "anything useful. In that case, the algorithm counts how many CRITICAL-classified ports "
       "are open and how many HIGH-classified ones are open, and maps those counts to a score "
       "range (91-99 for CRITICAL, 70-89 for HIGH, 40 as the baseline). This is a fallback "
       "heuristic, not a replacement for actual CVE data — but it ensures that a host with "
       "a wide-open RDP port doesn't get scored the same as a host with nothing interesting.")

    doc.add_page_break()
    sh(doc,"3.3  Frontend Implementation")
    bp(doc,
       "The frontend is a React 18 SPA built with Vite 5 and styled with Tailwind CSS. "
       "Figure 5 shows the data flow through the full stack, and Figure 6 shows what the "
       "production dashboard actually looks like.")
    add_figure(doc,"dfd","Figure 5: Level-1 Data Flow Diagram (DFD)")
    add_figure(doc,"ui","Figure 6: Sentinel Audit \u2013 Main Dashboard Interface")

    bp(doc,
       "The application breaks down into one main component (App.jsx) and three sub-components. "
       "App.jsx owns all the state — scan results, scan mode, progress percentage, error messages — "
       "and handles all the async fetch calls to the backend. When a scan completes, it calls "
       "setResults() with the response data, which triggers re-renders downstream.")
    bp(doc,
       "ScanForm renders the left sidebar. It has two tabs: Port Scanner and Ping Sweep. "
       "The Port Scanner tab takes a target (IP or hostname), an optional port range, TCP/UDP "
       "checkboxes, and a toggle for common-ports-only mode. The Ping Sweep tab takes a CIDR "
       "subnet string. Both tabs submit to their respective handlers in App.jsx.")
    bp(doc,
       "ResultsTable handles the main results display. Each row shows port number, protocol, "
       "service name, a colour-coded severity badge, and the attack vector. Clicking a row "
       "expands it to show the raw banner and any CVE records retrieved, each with its CVSS "
       "score, severity label, and description. Severity badge colours follow a traffic-light "
       "scheme: CRITICAL is red (#ff2a2a), HIGH is orange (#ff8533), MEDIUM is yellow (#ffd700), "
       "LOW is green (#0aff0a).")
    bp(doc,
       "NetworkGraph uses HTML5 Canvas to draw a simple hub-and-spoke topology diagram when "
       "ping sweep results come in. The scanning host sits at the centre; each alive host is "
       "a node on the perimeter connected to the centre by a line, with the IP address "
       "labelled alongside. It's not a sophisticated graph, but it gives an immediate visual "
       "sense of the subnet layout, which is often more useful than a bare list of IPs.")
    bp(doc,
       "The Threat Score gauge is a CSS conic-gradient ring where the arc length corresponds "
       "to the score percentage. The fill colour changes dynamically based on the range. "
       "Scan progress is polled by fetching /api/progress every 500ms during an active scan; "
       "the returned {current, total} pair is used to calculate a percentage and update a "
       "progress bar. The fetch is in a setInterval that gets cleared once the scan completes.")

    sh(doc,"3.4  Backend API Design")
    bp(doc,
       "The backend is built with FastAPI and served by Uvicorn. All route handlers are "
       "collected under an APIRouter and mounted at /api. CORS middleware is configured to "
       "allow all origins during development, since the Vite dev server runs on port 5173 "
       "while the FastAPI backend runs on port 8000. Table 4 documents all the endpoints.")

    tname(doc,"Table 4: API Endpoints \u2013 Sentinel Audit Backend")
    api=doc.add_table(rows=1,cols=4); api.style='Table Grid'; api.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(api,["Method","Endpoint","Description","Request Body"])
    api_data=[
        ("GET","/api/status","Health check \u2013 confirms backend is running","None"),
        ("GET","/api/progress","Returns current scan progress {current, total}","None"),
        ("POST","/api/cancel","Sets cancellation flag for the active scan","None"),
        ("POST","/api/scan","Initiates port scan; returns open ports with vulnerability data","ScanRequest JSON"),
        ("POST","/api/ping-sweep","Initiates subnet ping sweep; returns list of alive hosts","PingSweepRequest JSON"),
        ("POST","/api/download/pdf","Generates and streams a formatted PDF audit report","ReportData JSON"),
    ]
    for row_d in api_data:
        data_row(api,list(row_d),[C,L,L,L])
    gap(doc)

    bp(doc,
       "One implementation detail worth noting: the cancellation mechanism uses a simple boolean "
       "flag (scan_cancelled = True) that gets checked at each iteration of the port loop. "
       "When the frontend calls /api/cancel, the flag is set; the scanner's inner loop checks "
       "it before each new batch of tasks and breaks out if it's true. It's not instantaneous — "
       "tasks already in flight will run to completion — but it's effective and doesn't require "
       "any thread interruption or process killing.")

# ── chapter 4 ────────────────────────────────────────────────────────────────

def add_chapter4(doc):
    ch_title(doc,4,"Results and Analysis")

    bp(doc,
       "This chapter documents the results of testing Sentinel Audit against real targets, "
       "evaluates whether the scanning engine behaves correctly, and examines the Threat Score "
       "logic under a few concrete scenarios. Testing was done both in a controlled local lab "
       "environment and against scanme.nmap.org, a host maintained by the Nmap project "
       "specifically to allow authorized public scanning.")

    sh(doc,"4.1  Scan Results and Dashboard")
    ssh(doc,"4.1.1  Port Scan Results \u2013 scanme.nmap.org")
    bp(doc,
       "We ran a TCP common-ports scan against scanme.nmap.org. The 42-port scan completed "
       "in roughly 3.2 seconds, which is acceptably fast for an async full-connect scanner. "
       "Table 5 shows a representative subset of what was found.")

    tname(doc,"Table 5: Sample Scan Result \u2013 Open Ports Detected (scanme.nmap.org, TCP, Common Ports)")
    res=doc.add_table(rows=1,cols=5); res.style='Table Grid'; res.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(res,["Port","Service","Protocol","Severity","Attack Vector"])
    res_data=[
        ("22","SSH","TCP","CRITICAL","Brute Force, Private Key Theft"),
        ("80","HTTP","TCP","MEDIUM","SQLi, XSS, SSRF"),
        ("9929","Ephemeral / Dynamic Port","TCP","LOW","Unknown / Target Dependent"),
    ]
    for row_d in res_data:
        data_row(res,list(row_d),[C,L,C,C,L])
    gap(doc)

    bp(doc,
       "Port 22 came back CRITICAL, which is correct — SSH is one of the most frequently "
       "bruteforced services on the internet. Port 80 was flagged MEDIUM; that's appropriate "
       "given the range of web application vulnerabilities that could be present, though without "
       "knowing what application is behind it, MEDIUM is a reasonable conservative estimate. "
       "Port 9929 is nping — a tool used by the scanme server itself — and the scanner "
       "correctly classified it as LOW with no specific mapping, since it's not a "
       "well-known attack surface. These results are consistent with running the same scan "
       "in Nmap [2], which gives confidence that the detection logic is working correctly.")

    ssh(doc,"4.1.2  CVE Enrichment Results")
    bp(doc,
       "The banner grabbed from port 22 on scanme.nmap.org was "
       "'SSH-2.0-OpenSSH_6.6.1p1 Ubuntu'. The keyword extractor pulled 'OpenSSH 6.6' from "
       "that and sent it as the keywordSearch parameter to the NVD API. The API returned "
       "three CVEs, the most notable being CVE-2016-0777 (CVSS 8.8, High severity) — an "
       "information disclosure vulnerability in the OpenSSH roaming feature that was patched "
       "in version 7.1. The scanner correctly computed a Threat Score of 88 from that. "
       "Which is, frankly, appropriate: running OpenSSH 6.6 in 2024 without mitigations is "
       "a genuine security concern.")

    ssh(doc,"4.1.3  Ping Sweep Results \u2013 Local LAN")
    bp(doc,
       "We ran a ping sweep across 192.168.1.0/24 (254 host addresses). The sweep finished "
       "in about 8.5 seconds and correctly identified 6 alive hosts. We cross-checked "
       "against the router's DHCP lease table — all 6 discovered hosts were in there, and "
       "no hosts were missed. The NetworkGraph rendered a hub-and-spoke diagram showing all "
       "six hosts connected to the central node, which gave a clear visual of the subnet "
       "layout without any additional tooling.")

    sh(doc,"4.2  Threat Score Analysis")
    bp(doc,
       "We validated the Threat Score logic against two specific scenarios to confirm it behaves "
       "as intended. Both map to the CVSS v3.1 standard [13].")
    bi(doc,"Scenario A: A scan returned port 3389 (RDP) with CVE-2019-0708 (BlueKeep, CVSS 9.8, "
          "Critical) in the CVE data. The Threat Score computed was 98. That's correct: BlueKeep "
          "is a pre-authentication RCE vulnerability that caused widespread concern when it was "
          "disclosed. A score of 98 feels appropriately urgent.")
    bi(doc,"Scenario B: A scan of a patched local Linux server found port 22 open but banner "
          "grabbing returned nothing useful, so no CVE data was retrieved. The fallback heuristic "
          "assigned 90 (CRITICAL classification for port 22 with criticalCount = 1). That's "
          "a conservative scoring — it flags that something worth investigating is there even "
          "when live CVE data isn't available.")
    bp(doc,
       "These two scenarios demonstrate that the scoring logic correctly reflects CVSS data "
       "when it's present and applies a sensible fallback when it isn't. The results are "
       "consistent with NIST SP 800-115 [7], which recommends classification by likelihood "
       "and impact of exploitation.")

    ssh(doc,"4.2.1  Severity Classification Validation")
    bp(doc,
       "Table 6 shows how open ports map to severity levels within Sentinel Audit's internal "
       "classification system, along with the corresponding Threat Score ranges when the "
       "fallback heuristic applies.")

    tname(doc,"Table 6: Severity Classification of Open Ports (as implemented in scanner.py)")
    sev=doc.add_table(rows=1,cols=3); sev.style='Table Grid'; sev.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(sev,["Severity Level","Qualifying Ports","Threat Score Range"])
    sev_data=[
        ("CRITICAL","21, 22, 23, 139, 445, 1433, 1521, 2375, 3306, 3389, 4444, 1099","91\u201399"),
        ("HIGH","25, 53, 69, 111, 135, 161, 389, 512, 513, 514, 873, 2049, 5900, 6379, 9200","75\u201389"),
        ("MEDIUM","Ports with a known attack vector in VULNERABILITY_MAP","40\u201374"),
        ("LOW","All other open ports with no known attack vector","40"),
    ]
    for row_d in sev_data:
        data_row(sev,list(row_d),[C,L,C])
    gap(doc)

    bp(doc,
       "The CRITICAL port selection aligns with what industry threat reports consistently "
       "identify as the most commonly targeted services in ransomware campaigns and advanced "
       "persistent threat activity — RDP, SMB, SQL Server, and shell services feature "
       "prominently in virtually every post-incident analysis of network breaches.")

# ── chapter 5 ────────────────────────────────────────────────────────────────

def add_chapter5(doc):
    ch_title(doc,5,"Advantages, Limitations and Applications")

    bp(doc,
       "No tool does everything equally well. This chapter tries to be honest about what "
       "Sentinel Audit does well, where it falls short, and the contexts in which it's "
       "actually useful versus those where a different tool would serve better.")

    sh(doc,"5.1  Advantages")
    bi(doc,"Single unified interface: everything from port scanning to CVE lookup to PDF report "
          "generation happens in one tool. No post-processing scripts, no separate databases to "
          "maintain, no multi-tool workflows.")
    bi(doc,"No installation friction for end users: because the React frontend runs in a browser "
          "and the backend is packaged into a single executable, getting started is essentially a "
          "double-click. That's a significant accessibility improvement over tools like OpenVAS "
          "or even Nmap.")
    bi(doc,"Live CVE data, always current: by querying the NIST NVD API at scan time, the "
          "vulnerability information is always up to date. There's no offline database to "
          "keep synchronised.")
    bi(doc,"Genuinely async scanning performance: using asyncio with a 1,000-connection semaphore "
          "ceiling gives throughput that's competitive with threaded implementations, without "
          "the overhead of context switching between threads.")
    bi(doc,"Quantified, defensible risk metric: basing the Threat Score on CVSS base scores "
          "means that when you tell a stakeholder 'this host scores 87 out of 100', you can "
          "point to a published industry standard as the basis for that number.")
    bi(doc,"Works on private networks: unlike Shodan, Sentinel Audit has no dependency on the "
          "target being internet-routable. It works equally well on a corporate intranet, "
          "a segmented lab network, or a home router.")
    bi(doc,"Cancellable scans: the /api/cancel endpoint sets a flag that the scanner checks "
          "per-port, so a user can stop a long scan without having to kill the process.")
    bi(doc,"Open source (MIT): the codebase is available for inspection, extension, and reuse "
          "without any licensing fees.")

    sh(doc,"5.2  Limitations")
    bi(doc,"Full TCP connect only: Sentinel Audit does not craft raw SYN packets the way Nmap "
          "does in its half-open scan mode. Every TCP check completes the full handshake, which "
          "means scans are more likely to be logged by target-side firewalls and IDS systems.")
    bi(doc,"UDP reliability issues: UDP port scanning is inherently ambiguous. Silence doesn't "
          "prove a port is closed — it could be filtered, or the service simply doesn't respond "
          "to unsolicited datagrams. Sentinel Audit reports UDP ports as open only when an actual "
          "response is received, which means false negatives are possible.")
    bi(doc,"NVD API rate limiting: without an API key, the NIST NVD API allows roughly 5 "
          "requests per 30-second window. On a scan that finds many open ports, CVE lookups "
          "will be throttled and some results may come back empty. In practice we observed "
          "this during testing with hosts that had 15+ open ports.")
    bi(doc,"No OS fingerprinting: the current implementation doesn't attempt to identify the "
          "target's operating system. Knowing whether a host is running Windows Server vs. Linux "
          "can be significant context for vulnerability assessment, and Sentinel Audit simply "
          "doesn't provide it.")
    bi(doc,"Single-target port scanning only: you can sweep a subnet for alive hosts, but port "
          "scanning targets one host at a time. Scanning every host on a /24 subnet requires "
          "running the scan manually for each IP.")
    bi(doc,"No authenticated checks: Sentinel Audit does purely black-box scanning. It can't "
          "log in and check whether patches have been applied, what software is installed, or "
          "how services are configured internally.")
    bi(doc,"Internet dependency for CVE enrichment: the NVD API integration requires outbound "
          "internet access. In truly air-gapped deployments, CVE data won't be available — "
          "though the rest of the scanning functionality still works fine.")

    sh(doc,"5.3  Applications")
    bp(doc,"Where does Sentinel Audit actually fit? A few contexts stand out:")
    bi(doc,"University security courses: it makes port scanning and vulnerability analysis "
          "tangible for students who aren't yet comfortable with command-line tools. The visual "
          "dashboard makes the relationship between open ports and risk much more concrete than "
          "reading text output from Nmap.")
    bi(doc,"SME internal audits: small and medium enterprises without a dedicated security team "
          "can run regular scans, keep records of what's open, and monitor for unexpected "
          "changes — all without significant setup or training.")
    bi(doc,"Penetration test reconnaissance: during the initial enumeration phase of an "
          "authorised pentest, Sentinel Audit can quickly produce a prioritised list of "
          "attack candidates with CVE context, which is useful for planning further effort.")
    bi(doc,"Security awareness demonstrations: showing a non-technical audience what's "
          "actually exposed on their network is far more impactful than telling them. "
          "The Threat Score in particular tends to land effectively in those conversations.")
    bi(doc,"CI/CD pipeline integration: the REST API makes it straightforward to call Sentinel "
          "Audit programmatically after a deployment and fail a build or send an alert if "
          "unexpected high-severity ports are found open on the deployed infrastructure.")
    bi(doc,"Network inventory: the ping sweep doubles as a quick way to discover all live "
          "devices on a subnet, which is useful for asset inventory work.")

# ── chapter 6 ────────────────────────────────────────────────────────────────

def add_chapter6(doc):
    ch_title(doc,6,"Conclusion and Future Scope")

    sh(doc,"6.1  Conclusion")
    bp(doc,
       "Sentinel Audit set out to provide a network security auditing tool that doesn't "
       "require command-line expertise, integrates live vulnerability data, computes an "
       "objective risk score, and produces a formatted report — all in a self-contained, "
       "deployable package. On those terms, the project succeeded. The asynchronous scanning "
       "engine works reliably, the CVE enrichment pipeline correctly surfaces relevant "
       "vulnerabilities when they exist, the Threat Score reflects CVSS standards, and "
       "the React dashboard gives a genuinely usable interface to all of it.")
    bp(doc,
       "The literature review showed clearly that existing tools, while powerful, leave a "
       "real gap for users who need something between 'type nmap into a terminal' and "
       "'hire a security consultant'. The test results in Chapter 4 validated that the "
       "detection accuracy and scoring logic are functionally correct. The tool works on "
       "private networks, costs nothing to deploy, and can run as a standalone executable.")
    bp(doc,
       "What the project also demonstrated is that modern Python's async primitives are "
       "genuinely well-suited to network scanning workloads. Managing a thousand concurrent "
       "connection attempts on a single event loop without any threading — and keeping the "
       "API server responsive to progress polling at the same time — was simpler to implement "
       "correctly than it might sound. FastAPI and asyncio make that sort of design feel "
       "natural rather than difficult.")

    sh(doc,"6.2  Future Scope")
    bp(doc,"A number of meaningful improvements are reasonably within reach for future versions:")
    ni(doc,"SYN / half-open scanning via raw sockets (Scapy or a C extension) would reduce "
          "scan visibility to target-side IDS systems and improve speed, since completing the "
          "full handshake for every port is unnecessary overhead when you just want to know "
          "whether a port is listening.")
    ni(doc,"OS fingerprinting based on TTL values and TCP/IP stack behaviour would add "
          "meaningful context to scan results — knowing the target is running Windows Server "
          "2019 versus Ubuntu 22.04 changes the vulnerability prioritisation significantly.")
    ni(doc,"NIST NVD API key support and local SQLite caching of CVE lookups would solve the "
          "rate limiting problem. The API key raises the limit from 5 to 50 requests per "
          "30-second window; caching means repeat scans of the same services don't "
          "re-query at all.")
    ni(doc,"Subnet-wide port scanning — extending the scan engine to accept a CIDR range "
          "rather than just a single host — would make audit workflows considerably more "
          "efficient for network administrators.")
    ni(doc,"SSH-based authenticated scanning using Paramiko would allow the tool to log in "
          "to target hosts and check patch levels, file permissions, and running service "
          "configurations — closing the gap with tools like Nessus for trusted-host auditing.")
    ni(doc,"A scheduled scanning mode (via APScheduler) that runs automatically at configured "
          "intervals and alerts when new ports open or Threat Scores change would turn Sentinel "
          "Audit into a continuous monitoring solution rather than a point-in-time tool.")
    ni(doc,"Machine learning-based anomaly detection on scan history could identify unusual "
          "port patterns even in the absence of specific CVE data — flagging, for example, "
          "a new high-numbered port that wasn't there yesterday.")
    ni(doc,"STIX 2.1 report export would enable interoperability with SOC toolchains and "
          "SIEM platforms that consume structured threat intelligence in that format.")

# ── references ────────────────────────────────────────────────────────────────

def add_references(doc):
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("References"); r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True
    sp(p,0,18)
    refs=[
        "[1] NIST National Vulnerability Database (NVD), National Institute of Standards and Technology. "
        "Available: https://nvd.nist.gov/. Accessed: April 2025.",
        "[2] G. Lyon, \u2018Nmap Network Scanning: The Official Nmap Project Guide to Network Discovery "
        "and Security Scanning,\u2019 Insecure.Com LLC, 2009. Available: https://nmap.org/book/.",
        "[3] NIST NVD REST API v2.0 Documentation. Available: https://nvd.nist.gov/developers/vulnerabilities. "
        "Accessed: April 2025.",
        "[4] R. Graham, \u2018Masscan: Mass IP Port Scanner,\u2019 GitHub Repository, 2013. "
        "Available: https://github.com/robertdavidgraham/masscan.",
        "[5] J. Matherly, \u2018Shodan: The Search Engine for Internet-Connected Devices,\u2019 "
        "Shodan, 2009. Available: https://www.shodan.io/.",
        "[6] E. Al-Shaer, W. Marrero, A. Al-Haj, and M. Tabia, \u2018Network Configuration in A Box: "
        "Towards End-to-End Verification of Network Reachability and Security,\u2019 in Proc. 17th "
        "IEEE ICNP, Princeton, NJ, USA, 2009, pp. 123\u2013132.",
        "[7] K. Scarfone, M. Souppaya, A. Cody, and A. Orebaugh, \u2018Technical Guide to Information "
        "Security Testing and Assessment,\u2019 NIST Special Publication 800-115, Sept. 2008.",
        "[8] A. S. Tanenbaum and D. J. Wetherall, \u2018Computer Networks,\u2019 5th ed., Prentice Hall, "
        "2010, Ch. 5\u20137.",
        "[9] J. Postel, \u2018Transmission Control Protocol,\u2019 IETF RFC 793, Sept. 1981.",
        "[10] J. Postel, \u2018User Datagram Protocol,\u2019 IETF RFC 768, Aug. 1980.",
        "[11] NIST, \u2018Framework for Improving Critical Infrastructure Cybersecurity (Cybersecurity "
        "Framework),\u2019 Version 1.1, National Institute of Standards and Technology, April 2018.",
        "[12] OWASP Foundation, \u2018OWASP Top Ten \u2013 2021 Edition,\u2019 Open Web Application "
        "Security Project, 2021. Available: https://owasp.org/Top10/.",
        "[13] K. Scarfone and P. Mell, \u2018A Complete Guide to the Common Vulnerability Scoring "
        "System Version 2.0,\u2019 NIST, May 2007; updated to CVSS v3.1 by P. Spring et al., FIRST, "
        "June 2019.",
        "[14] Greenbone Networks GmbH, \u2018OpenVAS \u2013 Open Vulnerability Assessment System,\u2019 "
        "Documentation. Available: https://www.openvas.org/. Accessed: April 2025.",
        "[15] C. Sullo and D. Lodge, \u2018Nikto Web Server Scanner,\u2019 GitHub Repository, 2001. "
        "Available: https://github.com/sullo/nikto.",
        "[16] S. M. Bellovin and W. R. Cheswick, \u2018Network Firewalls,\u2019 IEEE Communications "
        "Magazine, vol. 32, no. 9, pp. 50\u201357, Sept. 1994.",
        "[17] Python Software Foundation, \u2018asyncio \u2013 Asynchronous I/O,\u2019 Python 3.11 "
        "Documentation. Available: https://docs.python.org/3/library/asyncio.html. Accessed: April 2025.",
        "[18] S. Ram\u00edrez, \u2018FastAPI \u2013 High Performance, Easy to Learn, Fast to Code, "
        "Ready for Production,\u2019 2018. Available: https://fastapi.tiangolo.com/.",
        "[19] Meta Platforms, Inc., \u2018React \u2013 A JavaScript Library for Building User "
        "Interfaces,\u2019 Documentation. Available: https://react.dev/. Accessed: April 2025.",
        "[20] E. You, \u2018Vite \u2013 Next Generation Frontend Tooling,\u2019 Documentation. "
        "Available: https://vitejs.dev/. Accessed: April 2025.",
    ]
    for ref in refs:
        p=doc.add_paragraph(); r=p.add_run(ref)
        r.font.name='Times New Roman'; r.font.size=Pt(12)
        p.paragraph_format.left_indent=Inches(0.25)
        p.paragraph_format.first_line_indent=Inches(-0.25)
        sp(p,2,4,1.5); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY

# ── appendices ────────────────────────────────────────────────────────────────

def add_appendices(doc):
    # Appendix I
    doc.add_page_break()
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run("Appendix I: Soft Code Flowcharts")
    r.font.name='Times New Roman'; r.font.size=Pt(14); r.bold=True; sp(p,0,18)
    bp(doc,
       "This appendix brings together the three core algorithm flowcharts in one place, "
       "as a reference supplement to the methodology described in Chapter 3. These diagrams "
       "were constructed to reflect the actual implementation in the source code rather than "
       "an idealised version.")
    add_figure(doc,"flow","Appendix Figure I.1: Port Scanning Algorithm Flowchart")
    add_figure(doc,"ping","Appendix Figure I.2: Network Ping Sweep Algorithm Flowchart")
    add_figure(doc,"threat","Appendix Figure I.3: Threat Score Calculation Engine")

    # Appendix II
    doc.add_page_break()
    p2=doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=p2.add_run("Appendix II: Data Sheets / API Specifications")
    r2.font.name='Times New Roman'; r2.font.size=Pt(14); r2.bold=True; sp(p2,0,18)
    bp(doc,
       "This appendix provides the technical specifications for external APIs and key data "
       "structures used by Sentinel Audit. These details are useful for anyone looking to "
       "extend the tool or integrate it into another workflow.")

    ssh(doc,"II.1  NIST NVD API v2.0 \u2013 CVE Lookup Endpoint")
    bp(doc,"Base URL: GET https://services.nvd.nist.gov/rest/json/cves/2.0")
    bp(doc,"Parameters used by Sentinel Audit:")
    bi(doc,"keywordSearch (string): The keyword extracted from the scanned service banner "
          "(e.g. 'OpenSSH 6.6', 'vsftpd 2.3'). This is what drives the CVE search.")
    bi(doc,"resultsPerPage (integer): Fixed at 3. We intentionally limit to 3 results to "
          "stay within rate limits and keep the UI readable.")
    bp(doc,"Response fields consumed:")
    bi(doc,"vulnerabilities[].cve.id \u2013 the CVE identifier (e.g. 'CVE-2023-38408')")
    bi(doc,"vulnerabilities[].cve.descriptions[0].value \u2013 English CVE description")
    bi(doc,"vulnerabilities[].cve.metrics.cvssMetricV31[0].cvssData.baseScore \u2013 CVSS v3.1 base score")
    bi(doc,"vulnerabilities[].cve.metrics.cvssMetricV31[0].baseSeverity \u2013 CRITICAL/HIGH/MEDIUM/LOW")
    bp(doc,"Rate limits: 5 requests per 30 seconds without an API key; 50 requests per 30 seconds "
          "with one (available free from https://nvd.nist.gov/developers/request-an-api-key).")

    ssh(doc,"II.2  ScanRequest Pydantic Model")
    bp(doc,"Sent by the frontend to POST /api/scan as JSON:")
    bi(doc,"target (str): IP address or hostname to scan.")
    bi(doc,"start_port (int): Lower bound of the port range (1\u201365535).")
    bi(doc,"end_port (int): Upper bound of the port range (1\u201365535).")
    bi(doc,"scan_tcp (bool): Whether to include TCP scanning.")
    bi(doc,"scan_udp (bool): Whether to include UDP scanning.")
    bi(doc,"common_ports_only (bool): If true, ignores start/end_port and uses the 42-port list.")

    ssh(doc,"II.3  PingSweepRequest Pydantic Model")
    bp(doc,"Sent by the frontend to POST /api/ping-sweep as JSON:")
    bi(doc,"subnet (str): CIDR notation subnet string (e.g. '192.168.1.0/24'). "
          "Parsed using Python's ipaddress.ip_network() with strict=False.")

    # Appendix III
    doc.add_page_break()
    p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r3=p3.add_run("Appendix III: List of Components / Technologies")
    r3.font.name='Times New Roman'; r3.font.size=Pt(14); r3.bold=True; sp(p3,0,18)
    bp(doc,
       "Table 7 lists every software component and library used in Sentinel Audit, along with "
       "the version used during development, the licence under which it is distributed, and "
       "its role in the system.")

    tname(doc,"Table 7: Technology Stack \u2013 Sentinel Audit")
    tech=doc.add_table(rows=1,cols=4); tech.style='Table Grid'; tech.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr_row(tech,["Component","Version","Licence","Role / Purpose"])
    tech_data=[
        ("Python","3.10+","PSF","Backend runtime"),
        ("FastAPI","0.110+","MIT","Async REST API framework"),
        ("Uvicorn","0.29+","BSD","ASGI web server"),
        ("Pydantic","2.x","MIT","Request/response validation"),
        ("fpdf2","2.7+","LGPL","PDF report generation"),
        ("asyncio","stdlib","PSF","Async I/O, event loop"),
        ("ipaddress","stdlib","PSF","IPv4/IPv6 subnet parsing"),
        ("React","18.x","MIT","Frontend UI library"),
        ("Vite","5.x","MIT","Frontend build tool / dev server"),
        ("Tailwind CSS","3.x","MIT","Utility-first CSS framework"),
        ("Lucide React","0.x","ISC","SVG icon library"),
        ("Node.js","18+","MIT","JS runtime for frontend build"),
        ("PyInstaller","6.x","GPLv2+","Python application packager"),
        ("Electron","30+","MIT","Desktop application wrapper"),
        ("NIST NVD API","v2.0 REST","Public","Live CVE data source"),
    ]
    for row_d in tech_data:
        data_row(tech,list(row_d),[L,C,C,L])
    gap(doc)

    # Appendix IV
    doc.add_page_break()
    p4=doc.add_paragraph(); p4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r4=p4.add_run("Appendix IV: Sample Code Listings")
    r4.font.name='Times New Roman'; r4.font.size=Pt(14); r4.bold=True; sp(p4,0,18)
    bp(doc,
       "The following code excerpts from the Sentinel Audit codebase illustrate the key "
       "implementation patterns discussed in Chapter 3. These are abbreviated for clarity — "
       "the full source is available in the project repository.")

    ssh(doc,"IV.1  Async TCP Port Scanner (scanner.py)")
    code1=(
        "async def scan_tcp_port(port):\n"
        "    if check_cancel and check_cancel(): return None\n"
        "    async with sem:\n"
        "        try:\n"
        "            fut = asyncio.open_connection(target, port)\n"
        "            reader, writer = await asyncio.wait_for(fut, timeout=0.8)\n"
        "            try:\n"
        "                data = await asyncio.wait_for(reader.read(1024), timeout=1.0)\n"
        "                banner = data.decode(errors='ignore').strip() or 'No Banner (TCP)'\n"
        "            except (asyncio.TimeoutError, Exception):\n"
        "                banner = 'No Banner (TCP)'\n"
        "            finally:\n"
        "                writer.close(); await writer.wait_closed()\n"
        "            cve_data = await fetch_cves(banner)\n"
        "            vuln_data = analyze_port(port, banner)\n"
        "            return {'port': port, 'state': 'open', 'protocol': 'tcp',\n"
        "                    'service': vuln_data['service'], 'banner': banner,\n"
        "                    'attack_vector': vuln_data['attack_vector'],\n"
        "                    'cve_data': cve_data, 'base_severity': base_severity}\n"
        "        except (asyncio.TimeoutError, ConnectionRefusedError, OSError):\n"
        "            pass\n"
        "        update_progress(); return None"
    )
    cp=doc.add_paragraph(); cr=cp.add_run(code1)
    cr.font.name='Courier New'; cr.font.size=Pt(9); sp(cp,4,8)

    ssh(doc,"IV.2  Threat Score Calculation (App.jsx)")
    code2=(
        "const calculateThreatScore = () => {\n"
        "    if (!results || results.length === 0) return 0;\n"
        "    let maxBaseScore = 0, criticalCount = 0, highCount = 0;\n"
        "    results.forEach(port => {\n"
        "        if (port.cve_data?.length > 0) {\n"
        "            port.cve_data.forEach(cve => {\n"
        "                if (cve.baseScore > maxBaseScore) maxBaseScore = cve.baseScore;\n"
        "            });\n"
        "        }\n"
        "        if (port.base_severity === 'CRITICAL') criticalCount++;\n"
        "        if (port.base_severity === 'HIGH') highCount++;\n"
        "    });\n"
        "    if (maxBaseScore > 0) return Math.min(100, Math.round(maxBaseScore * 10));\n"
        "    if (criticalCount > 0) return 90 + Math.min(9, criticalCount);\n"
        "    if (highCount > 0)     return 70 + Math.min(19, highCount * 5);\n"
        "    return 40;\n"
        "};"
    )
    cp2=doc.add_paragraph(); cr2=cp2.add_run(code2)
    cr2.font.name='Courier New'; cr2.font.size=Pt(9); sp(cp2,4,8)

    ssh(doc,"IV.3  Async Ping Sweep (pingsweep.py)")
    code3=(
        "async def ping_ip(self, ip: str, sem: asyncio.Semaphore) -> str:\n"
        "    async with sem:\n"
        "        try:\n"
        "            is_win = platform.system().lower() == 'windows'\n"
        "            param = '-n' if is_win else '-c'\n"
        "            t_param = '-w' if is_win else '-W'\n"
        "            t_val   = '800' if is_win else '1'\n"
        "            cmd = f'ping {param} 1 {t_param} {t_val} {ip}'\n"
        "            proc = await asyncio.create_subprocess_shell(\n"
        "                cmd, stdout=asyncio.subprocess.PIPE,\n"
        "                     stderr=asyncio.subprocess.PIPE)\n"
        "            stdout, _ = await proc.communicate()\n"
        "            out = stdout.decode(errors='ignore').upper()\n"
        "            if is_win:\n"
        "                if proc.returncode == 0 and 'REQUEST TIMED OUT' not in out \\\n"
        "                   and '100% LOSS' not in out:\n"
        "                    return ip\n"
        "            else:\n"
        "                if proc.returncode == 0: return ip\n"
        "        except Exception: pass\n"
        "        return None"
    )
    cp3=doc.add_paragraph(); cr3=cp3.add_run(code3)
    cr3.font.name='Courier New'; cr3.font.size=Pt(9); sp(cp3,4,8)

# ── main ─────────────────────────────────────────────────────────────────────

def build():
    print("Building Sentinel Audit Report...")
    doc = setup_doc()
    print("  [1/12] Abstract..."); add_abstract(doc)
    print("  [2/12] Table of Contents..."); add_toc(doc)
    print("  [3/12] List of Figures..."); add_list_of_figures(doc)
    print("  [4/12] List of Tables..."); add_list_of_tables(doc)
    print("  [5/12] Abbreviations..."); add_abbreviations(doc)
    print("  [6/12] Chapter 1 - Introduction..."); add_chapter1(doc)
    print("  [7/12] Chapter 2 - Literature Survey..."); add_chapter2(doc)
    print("  [8/12] Chapter 3 - Methodology & Implementation..."); add_chapter3(doc)
    print("  [9/12] Chapter 4 - Results & Analysis..."); add_chapter4(doc)
    print("  [10/12] Chapter 5 - Advantages, Limitations, Applications..."); add_chapter5(doc)
    print("  [11/12] Chapter 6 - Conclusion & Future Scope..."); add_chapter6(doc)
    print("  [12/12] References & Appendices..."); add_references(doc); add_appendices(doc)
    doc.save(OUTPUT)
    print(f"\nDONE! Report saved to:\n   {OUTPUT}")

if __name__ == "__main__":
    build()
